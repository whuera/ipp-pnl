from docx import Document 
from datetime import date
from docxtpl import DocxTemplate

# Crear un documento Word
doc = Document()

# Título del ensayo
doc.add_heading('Carta de Motivación', level=1)

# Fecha
today = date.today()
doc.add_paragraph(f"Fecha: {today.strftime('%d/%m/%Y')}")

# Información personal
doc.add_paragraph("Nombre: Juan Pérez\nTeléfono: +56 9 1234 5678\nCorreo Electrónico: juan.perez@example.com")

# Primer párrafo
doc.add_heading('Primer Párrafo', level=2)
doc.add_paragraph(
    "Estimados miembros del comité de admisiones,\n\n"
    "Me dirijo a ustedes con el deseo de postularme para la maestría en el programa académico modalidad online. "
    "Estoy muy entusiasmado con la posibilidad de formar parte de este prestigioso programa y considero que "
    "mi formación y experiencias previas me convierten en un candidato ideal para ello."
)

# Segundo párrafo
doc.add_heading('Segundo Párrafo', level=2)
doc.add_paragraph(
    "A lo largo de mi carrera, he desarrollado habilidades y conocimientos que considero fundamentales para "
    "mi éxito en esta maestría. Poseo un título en Ingeniería de Sistemas, y he trabajado durante cinco años "
    "en el área de desarrollo de software, donde he adquirido competencias en programación, gestión de proyectos "
    "y trabajo en equipo. Además, he participado en diversos cursos y seminarios relacionados con la tecnología "
    "y la innovación, lo que me ha permitido mantenerme actualizado en mi campo."
)

# Tercer párrafo
doc.add_heading('Tercer Párrafo', level=2)
doc.add_paragraph(
    "Mi intención es continuar mi desarrollo académico y profesional a través de esta maestría, y estoy convencido "
    "de que el programa online me proporcionará la flexibilidad y los recursos necesarios para lograr mis objetivos. "
    "Espero poder aportar mis conocimientos y experiencias al programa, así como aprender de los distinguidos profesores "
    "y compañeros que forman parte de él. Estoy plenamente comprometido a aprovechar al máximo esta oportunidad y "
    "a contribuir positivamente al entorno académico."
)

# Firma
doc.add_heading('Firma', level=2)
doc.add_paragraph("Atentamente,\n\nJuan Pérez")

# Guardar el documento
file_path = "Carta_de_Motivacion.docx"

doc.save(file_path)

file_path
